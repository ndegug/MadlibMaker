import tkinter as tk #main GUI tools
from tkinter import scrolledtext, messagebox
import webbrowser #for opening html files in browser
import tempfile #for WORD input file analysis
import re #regular expression library for substituting words
import os #for file and folder writing and reading
from long_strings_gui import * #collection of long strings
from docx import Document #for writing and reading word docs
from docx.shared import Pt #for word doc formatting
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #for word doc formatting

numword_dic = {}
unnumbered = "(/...)"
numbered = "(/...[0-9]+)"
customreg = "(/ct[0-9]+)"
numcustreg = "(/ct[0-9]+_[0-9]+)"
tail = "(_[0-9])"
htmlsample = '<span class="nowrap" style="display: none; display: inline-block; vertical-align: top; text-align: ' \
             'center;"><span style="display: block; padding: 0 0.2em;">__________</span><span style="display: block; ' \
             'font-size: 70%; line-height: 1em; padding: 0 0.2em;"><span style="position: relative; line-height: 1em; ' \
             'margin-top: -.2em; top: -.2em;">underscript</span></span></span>'
class MadlibApp:
    def __init__(self, root):
        self.root = root
        self.reset()
    def reset(self):
        self.custom = {} #full custom dictionary
        self.custom_keys = [] #storage for recording custom keys detected
        self.custom_index = 0 #for custom detection
        self.save_flag = False #flag for saving numbered keywords
        self.folders() #creates input and output folders if not here
        self.outlist = [] #output array for filled madlibs
        self.htlist = [] #output array for html madlibs
        self.title = '' #madlib title
        self.load_mode = False  # decides mode (False=write or True=load)
        self.welcomeMenuHandler() #welcome menu
    def load_input_file(self):
        # Clear existing widgets if necessary
        for widget in self.root.winfo_children():
            widget.destroy()
        self.load_mode = True #begin load mode
        label = tk.Label(self.root, text="Select a file to load:", font=("Arial", 12, "bold"),width=100, height=5, bg="#d0e7ff",
                         fg="black")
        label.pack(pady=10)

        inputs_path = os.path.join(os.getcwd(), "inputs")

        if not os.path.exists(inputs_path): #double check for inputs folder todo: add this check for every file write operation
            os.makedirs(inputs_path)

        files = [f for f in os.listdir(inputs_path) if f.endswith('.txt') or f.endswith('.docx')] #grabs all files in inputs

        if not files: #edge case of no valid input files
            no_files_label = tk.Label(self.root, text="No input files found in the 'inputs' folder.",
                                      font=("Arial", 12))
            no_files_label.pack(pady=10)
            return

        button_frame = tk.Frame(self.root)
        button_frame.pack(pady=10)

        row = 0  # initializes row and column counters for button grid
        col = 0
        for filename in files: #generate buttons for all files
            full_path = os.path.join(inputs_path, filename)
            btn = tk.Button(button_frame, text=filename, font=("Arial", 12),
                            command=lambda path=full_path: self.process_input_file_3(path), bg="#3b9dd3",
                        fg="white")
            btn.grid(row=row, column=col, padx=2, pady=2,
                     sticky="ew")  # defines the button's location on the grid ("ew" centers all buttons to their grid position)
            col += 1
            if col >= 5:  # max number of columns, then new row
                col = 0
                row += 1

    def process_input_file_3(self, path): #todo: comment lines
        custom_dict = {} #temporary custom dic storage
        title_text = "" #temporary title storage
        file_path = path
        try:
            # Read file content
            if file_path.endswith(".txt"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_path.endswith(".docx"):
                doc = Document(file_path)
                content = "\n".join([p.text for p in doc.paragraphs])
            else:
                messagebox.showerror("Unsupported Format", "Only .txt and .docx files are supported.")
                return

            # Extract title if present
            if "<t>" in content and "</t>" in content:
                start = content.index("<t>") + len("<t>")
                end = content.index("</t>")
                title_text = content[start:end].strip()

                # Remove title from content
                content = content[:content.index("<t>")] + content[end + len("</t>"):]

            # Extract custom dictionary
            if "<C>" in content:
                c_index = content.index("<C>")
                madlib_text = content[:c_index].strip()
                dict_text = content[c_index + len("<C>"):].strip()
                try:
                    custom_dict = eval(dict_text)
                except Exception as e:
                    messagebox.showerror("Custom Dict Error", f"Failed to parse custom dictionary:\n{e}")
                    return
            else:
                madlib_text = content.strip()

            self.title = title_text
            self.custom = custom_dict
            self.raw_in=madlib_text
            self.advance_from_first()

        except Exception as e:
            messagebox.showerror("File Error", f"Could not read file:\n{e}")

    def folders(self): #creates input and output folders if none exist
        if not os.path.isdir(os.path.join(os.getcwd(), "inputs")):
            os.mkdir(os.path.join(os.getcwd(), "inputs"))
        if not os.path.isdir(os.path.join(os.getcwd(), "outputs")):
            os.mkdir(os.path.join(os.getcwd(), "outputs"))
    def dummyscreen(self,name): #placeholder for menus-to-be
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        text = "This is the screen for " + str(name)
        w = tk.Label(self.root, text=text, width=80, height=10, bg="#d0e7ff", fg="black")
        w.pack(pady=10)
        self.root.mainloop()  # deploys the GUI screen till closed

    def file_write(self, content, name_of_file, path, ext):
        completeName = os.path.join(path, name_of_file + ext) #filepath with name and extension
        f = open(completeName, "w") #open file with path
        f.write(content) #insert contents
        f.close() #close the file

    def welcomeMenuHandler(self): #Welcome menu
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets

        # welcome text
        w = tk.Label(self.root, text='Hello!\n Welcome to the Madlib Maker',font=("Arial", 12, "bold"), width=80, height=10, bg="#d0e7ff", fg="black")
        w.pack(pady=(15))
        #smaller label to prompt choice
        w = tk.Label(self.root, text='What would you like to do\ntoday?', width=24,
                     height=3, bg="#d0e7ff", fg="black")
        w.pack(pady=10)

        # buttons for Welcome menu selection
        button_frame = tk.Frame(self.root)  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        # manual input button
        btn = tk.Button(button_frame, command=lambda: self.setup_first_window(), text="Manual Input", bg="#3b9dd3",
                        fg="white")  # defines each button with frame
        btn.grid(row=1, column=0, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid ("ew" centers all buttons to their grid position)
        # Load inputs button
        btn = tk.Button(button_frame, command=lambda: self.load_input_file(), text="Load a Madlib",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame
        btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")
        # Instruction menu button
        btn = tk.Button(button_frame, command=lambda: self.instruct_main(), text="Instructions",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame
        btn.grid(row=1, column=3, padx=2, pady=2,
                 sticky="ew")
        self.root.mainloop()  # deploys the GUI screen till closed
    def instruct_main(self): #Instructions top menu
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets

        # Instructions text
        w = tk.Label(self.root, text='What would you like to learn about?', font=("Arial", 12, "bold"), width=80, height=10, bg="#d0e7ff", fg="black")

        w.pack(pady=10)
        # buttons for Welcome menu selection
        button_frame = tk.Frame(self.root)  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        # How to play instruct
        btn = tk.Button(button_frame, command=lambda: self.end_instructions(how_to_play_madlibs, False), text="How to play Madlibs", bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=0, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # How to write instruct button
        btn = tk.Button(button_frame, command=lambda: self.write_instructions_menu(), text="How to write madlibs",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # Loading madlib instruct button
        btn = tk.Button(button_frame, command=lambda: self.end_instructions(loading_a_madlib,False), text="Loading a madlib",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=3, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # printing madlib instruct button
        btn = tk.Button(button_frame, command=lambda: self.end_instructions(printing_madlibs,False), text="Printing madlibs",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=4, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        #back to menu from instruct 1
        btn= tk.Button(self.root, text="< Back to main menu", command=lambda: self.reset(), bg="#3b9dd3",
                                    fg="white")
        btn.pack(pady=10)
        self.root.mainloop()  # deploys the GUI screen till closed

    def write_instructions_menu(self):
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        self.display = scrolledtext.ScrolledText(self.root, width=80, height=10, font=("Arial", 12), bg="#9cc9e0",
                                                 fg="black", wrap=tk.WORD)
        self.display.pack(pady=10, padx=5) #display widget spacing, horizontal and vertical
        self.display.insert(tk.END, how_to_write_a_madlib) #inserts "write madlibs" basic tutorial
        self.display.config(state='disabled')  # Make it read-only
        # buttons frame
        button_frame = tk.Frame(self.root)  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        # Generic words
        btn = tk.Button(button_frame, command=lambda: self.end_instructions(generic_words_list,True), text="Generic words",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame
        btn.grid(row=1, column=0, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        #numbered words
        btn = tk.Button(button_frame, command=lambda: self.end_instructions(numbered_words,True),
                        text="Numbered words",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=1, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # Custom words
        btn = tk.Button(button_frame, command=lambda: self.end_instructions(custom_words_basics,True),
                        text="Custom words basics",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # numbered custom words
        btn = tk.Button(button_frame, command=lambda: self.end_instructions(numbered_custom_words,True),
                        text="Numbered custom words",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=3, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # Custom dic
        btn = tk.Button(button_frame, command=lambda: self.end_instructions(prewriting_custom_configurations,True),
                        text="Preconfiguring custom words",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=4, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # back button

        btn = tk.Button(button_frame, command=lambda: self.instruct_main(), text="< Back to Instructions",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=2, column=1, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        #back to menu button
        btn = tk.Button(button_frame, command=lambda: self.reset(), text="<< Back to main menu",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=2, column=3, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
    def end_instructions(self,content,lv):#generic window for instructions at the end of the instruction tree.
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        self.display = scrolledtext.ScrolledText(self.root, width=80, height=10, font=("Arial", 12), bg="#9cc9e0",
                                                 fg="black", wrap=tk.WORD)
        self.display.pack(pady=10, padx=5)
        self.display.insert(tk.END, content)
        self.display.config(state='disabled')  # Make it read-only

        button_frame = tk.Frame(self.root)  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames

        if lv: #second level of menus, currently only used in specifics of writing madlibs, change to an int if more are added
            # back to write instructions button
            btn = tk.Button(button_frame, command=lambda: self.write_instructions_menu(), text="< Back to writing madlibs",
                            bg="#3b9dd3", fg="white")
            btn.grid(row=1, column=1, padx=2, pady=2,
                     sticky="ew")
            #back to instructions from lv2 end instructions
            btn = tk.Button(button_frame, command=lambda: self.instruct_main(), text="<< Back to Instructions",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame
            btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
            #back to menu from lv2 end instructions
            btn = tk.Button(button_frame, command=lambda: self.reset(), text="<<< Back to main menu",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
            btn.grid(row=1, column=3, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        else: #normal end instructions, only back to menu and back to instructions
            btn = tk.Button(button_frame, command=lambda: self.instruct_main(), text="< Back to Instructions",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
            btn.grid(row=1, column=1, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
            btn = tk.Button(button_frame, command=lambda: self.reset(), text="<< Back to main menu",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
            btn.grid(row=1, column=3, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
    def setup_first_window(self): #setup manual input window
        for widget in self.root.winfo_children():
            widget.destroy()
        l = tk.Label(self.root, text='Type your Madlib below!', font=("Arial", 12, "bold"), fg="black")
        l.pack(pady=5)


        self.input_entry = tk.Entry(self.root, font=("Arial", 14), width=80, bg="#d0e7ff", fg="black") #defines the text input field, size, color of font and background
        self.input_entry.pack(pady=10)# sets the verticle spacing given between the input field and other successive elements
        self.input_entry.focus_set()  # This sets focus so the cursor appears in the field
        l = tk.Label(self.root, text='Use these buttons to quick-drop a keyword!', font=("Arial", 12), fg="black")
        l.pack(pady=5)
        button_frame = tk.Frame(self.root) #defines the button frame
        button_frame.pack(pady=5)

        row = 0 #initializes row and column counters for button grid
        col = 0
        for key, label in generic_words.items(): #for every generic word. Grabs the key (/adj) and label (adjective)
            btn = tk.Button(button_frame, text=label, command=lambda k=key: self.insert_keyword(k), bg="#3b9dd3", fg="white") #defines each button with frame
            btn.grid(row=row, column=col, padx=2, pady=2, sticky="ew") #defines the button's location on the grid
            col += 1
            if col >= 10: #ten columns, then new row
                col = 0
                row += 1
        #define submit button, command to advance to next screen
        self.submit_btn = tk.Button(self.root, text="Submit", command=self.advance_from_first, bg="#3b9dd3", fg="white")
        self.submit_btn.pack(pady=10)
        #define display
        self.display = scrolledtext.ScrolledText(self.root, width=80, height=10, font=("Arial", 12), bg="#9cc9e0", fg="black")
        self.display.pack(pady=10)

        self.input_entry.bind("<KeyRelease>", self.sync_display) #syncs display on every key release

    def sync_display(self, event=None): #syncs text between manual entry and display
        self.display.delete(1.0, tk.END) #clears the display before updating with current text (from row 1 char 0 to the end)
        self.display.insert(tk.END, self.input_entry.get()) #inserts text from the end (which is the start of the field after delete) to
    def insert_keyword(self, keyword):#inserts a keyword into active input field
        self.input_entry.insert(tk.END, keyword) #inputs keyword without a space (do not have both lines active)
        self.sync_display() #adds the keyword to the display
    def advance_from_first(self): #handles input data
        # grabs manual input if manual
        if self.load_mode == False:#manual input
            self.raw_in = self.input_entry.get()
        elif self.load_mode == True: #from file
            pass
        else:
            self.dummyscreen('Invalid mode for advance_from_first()')

        self.userMadlib= self.raw_in.split(' ') #splits madlib body into array by spaces
        self.prompt_words = iter(self.userMadlib) #sets up iteration

        # Extract and sort custom keys
        all_custom_matches = re.findall(r'/ct(\d+)(?:_\d+)?', self.raw_in)
        self.custom_keys = sorted(set(f"/ct{id}" for id in all_custom_matches), key=lambda x: int(x[3:]))

        if self.custom_keys and self.custom=={}:#if customs detected
            self.custom_configure_window() #configure customs, file_choice will be executed as well
        else: #if no others are needed
            self.title_check()

    def title_check(self): #check if title is set
        if self.title:
            self.file_choice() #advance to input preview
        else:
            self.title_write() #write title
    def title_write(self): #write title
        for widget in self.root.winfo_children():
            widget.destroy()
        tk.Label(self.root, text="Would you like to title your Madlib?\n Type it here and click enter or \"skip\" to skip this step.", font=("Arial", 12, "bold")).pack()
        self.input_entry = tk.Entry(self.root, font=("Arial", 14), width=80, bg="#d0e7ff", fg="black")
        self.input_entry.bind("<Return>",
                              lambda event: self.title_saver()) #"enter" key submits title
        self.input_entry.pack(pady=10)
        self.input_entry.focus_set()  # automatically puts the cursor into the entry field
        button_frame = tk.Frame(self.root)  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        # Submit button to save title
        btn = tk.Button(button_frame, command=lambda: self.title_saver(), text="Submit", bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=0, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # Skip button, go to input preview
        btn = tk.Button(button_frame, command=lambda: self.file_choice(), text="Skip",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")

    def title_saver(self): #saves title
        self.title = self.input_entry.get().strip()
        self.file_choice()
    def custom_configure_window(self): #custom word configure window
        for widget in self.root.winfo_children():
            widget.destroy()

        self.display = scrolledtext.ScrolledText(self.root, width=80, height=10, font=("Arial", 12), bg="#9cc9e0", fg="black")
        self.display.pack(pady=10)
        self.display.insert(tk.END, "Custom words detected, please configure them.\n")

        self.custom_entry = tk.Entry(self.root, width=40,font=("Arial", 12), bg="#d0e7ff", fg="black")
        self.custom_entry.bind("<Return>",
                               lambda event: self.save_custom_word())  # allows the "enter" key to submit the custom word
        self.custom_entry.pack(pady=10)
        self.custom_button = tk.Button(self.root, text="Enter",bg="#3b9dd3", fg="white", command=self.save_custom_word)
        self.custom_button.pack(pady=5)

        self.prompt_next_custom() #prompt the next custom word in the display
        self.root.mainloop()  # deploys the GUI screen till closed

    def prompt_next_custom(self): #promts each custom word in display
        if self.custom_index < len(self.custom_keys):
            current_key = self.custom_keys[self.custom_index] #grab the custom key based on sequential index
            self.display.insert(tk.END, f"Custom {current_key[3:]}: ") #inserts each prompt with the entry based on ID number
        else:
            self.title_check() #with all customs configured, advance to title check


    def save_custom_word(self): #saves custom word to custom dictionary
        current_key = self.custom_keys[self.custom_index]
        user_input = self.custom_entry.get().strip()
        if user_input:
            self.custom[current_key] = user_input #adds user input to custom dictionary for current key
            self.custom_entry.delete(0, tk.END) #deletes user input from the input field
            self.custom_index += 1 #updates the custom keys index
            self.display.insert(tk.END, f"{user_input}\n") #append user input to display for reference todo: add this to main word entries
            self.prompt_next_custom() #prompt the next custom word to be configured

    def second_window(self): #generates Madlib play/fill window
        for widget in self.root.winfo_children():widget.destroy()
        tk.Label(self.root, text="Give us a/an:", font=("Arial", 12, "bold")).pack()

        self.display = scrolledtext.ScrolledText(self.root, width=80, height=10, font=("Arial", 12), bg="#9cc9e0",
                                                 fg="black")
        self.display.pack(pady=10, padx=5)
        tk.Label(self.root, text="Your entry:", font=("Arial", 12)).pack()
        # define and create entry field for user's entry for a word
        self.input_entry = tk.Entry(self.root, font=("Arial", 14), width=40, bg="#d0e7ff", fg="black")
        self.input_entry.bind("<Return>",
                              lambda event: self.process_next_keyword())  # allows the "enter" key to submit the keyword
        self.input_entry.pack(pady=10)
        self.input_entry.focus_set()  # automatically puts the cursor into the entry field

        self.submit_btn = tk.Button(self.root, text="Submit", command=self.process_next_keyword, bg="#3b9dd3",
                                    fg="white")
        self.submit_btn.pack(pady=10)

    def next_prompt(self): #iterates through each word of the madlib array, returns when user input is needed todo: complete comments
        try:
            self.current_word = next(self.prompt_words)

            #Generic word unnumbered
            if re.findall(unnumbered, self.current_word) and not re.findall(numbered,self.current_word) and not re.findall(customreg, self.current_word):
                regkey = re.findall(unnumbered, self.current_word)
                realkey = ''.join(regkey)
                self.save_key = realkey
                if realkey in generic_words:
                    self.display.insert(tk.END, f"{generic_words[realkey]}: ")
                    return
                elif realkey in ignored_words:
                    return
                else:
                    self.display.insert(tk.END,
                                        f"{realkey} is not a valid keyword, what would you like to replace it with?: ")
                    return

            elif re.findall(customreg, self.current_word) and not re.findall(numcustreg, self.current_word):
                # custom
                regkey = re.findall(customreg, self.current_word)
                realkey = ''.join(regkey)
                self.save_key = realkey
                if realkey in self.custom:
                    self.display.insert(tk.END, f"{self.custom[realkey]}: ")
                    return
                elif realkey not in self.custom:
                    self.display.insert(tk.END,f"{realkey} hasn't been configured, what would you like to replace it with?: ")
                    return

            elif re.findall(numcustreg, self.current_word):
                # numbered custom

                regkey = re.findall(numcustreg, self.current_word)
                realkey = ''.join(regkey)

                base = re.findall(customreg, self.current_word)
                base = ''.join(base)
                self.save_key = realkey
                if realkey not in numword_dic:# todo: fix unconfigged numbered customs, try adding "and base in self.custom"
                    # numbered cust unsaved
                    base = re.findall(customreg, self.current_word)
                    base = ''.join(base)
                    self.display.insert(tk.END, f"{self.custom[base]} : ")
                    self.save_flag = True
                    return
                elif realkey in numword_dic:
                    self.current_word = re.sub(realkey, numword_dic[realkey], self.current_word)
                    # numbered cust saved
                elif base not in self.custom:
                    # numbered cust unsaved
                    self.display.insert(tk.END,
                                        f"{realkey} hasn't been configured, what would you like to replace it with?: ")
                    self.save_flag = True
                    return
                else:
                    # others
                    self.display.insert(tk.END,
                                        f"{realkey} is not a valid keyword, what would you like to replace it with?: ")
                    return

            elif re.findall(numbered, self.current_word):
                # numbered
                regkey = re.findall(numbered, self.current_word)
                regkeyb = re.findall(unnumbered, self.current_word)
                realkey = ''.join(regkey)
                base = ''.join(regkeyb)
                self.save_key = realkey
                if realkey not in numword_dic.keys() and base in generic_words.keys():
                    self.display.insert(tk.END, f"{generic_words[base]}: ")
                    self.save_flag = True
                    return


                elif realkey not in numword_dic.keys() and base not in generic_words.keys():
                    # numbered unsaved

                    self.display.insert(tk.END,f"{realkey} is not a valid keyword, what would you like to replace it with?: ")
                    return
                elif realkey in numword_dic.keys():
                    self.current_word =  re.sub(realkey, numword_dic[realkey], self.current_word)
            else:
                pass
            self.outlist.append(self.current_word)
            self.next_prompt()
        except StopIteration:
            self.third_window()

    def advance_to_second(self): #advance to second window, start word prompting when selection is made
        self.second_window()
        self.next_prompt()

    def spoiler(self): #inserts madlib input text to display if the user chooses to spoil a loaded madlib
        self.display.insert(tk.END, "\n\n" + self.title + "\n\n"+self.raw_in)
        self.spbtn.destroy()

    def file_choice(self):
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        self.display = scrolledtext.ScrolledText(self.root, width=80, height=20, font=("Arial", 12), bg="#9cc9e0", fg="black", wrap=tk.WORD)
        self.display.pack(pady=20)
        w = tk.Label(self.root, text='What would you like to do with it?', width=40, height=5, bg="#d0e7ff", fg="black")
        w.pack(pady=5)
        button_frame = tk.Frame(self.root)  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames

        if not self.load_mode: #manual input, show it, no hide needed
            self.display.insert(tk.END, "\nHere is your Madlib:\n\n" + self.raw_in)
        else: #load, use button to reveal
            self.display.insert(tk.END, "\n   Your Madlib is ready to play. If you are the author, click \"SPOIL\" to preview it. Otherwise, we recommend you play it blind.")
            self.spbtn = tk.Button(button_frame, command=lambda: self.spoiler(), text="SPOIL",
                            bg="#F23F3F",
                            fg="white")  # defines each button with frame,
            self.spbtn.grid(row=1, column=4, padx=2, pady=2,
                     sticky="ew")


        #plain input save
        btn = tk.Button(button_frame, command=lambda: self.output_file_name(3), text="Save plain text input", bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=0, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        btn = tk.Button(button_frame, command=lambda: self.output_file_name(4), text="Save Word doc input",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=1, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # play without saving button
        btn = tk.Button(button_frame, command=lambda: self.advance_to_second(), text="Play without saving", bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        btn = tk.Button(button_frame, command=lambda: self.advance_to_html(), text="Print Physical",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=3, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        self.root.mainloop()  # deploys the GUI screen till closed
    def process_next_keyword(self):
        user_text = self.input_entry.get().strip() #grab user text
        self.display.insert(tk.END, f"{user_text}\n") #display next to prompt in display
        self.input_entry.delete(0, tk.END) #delete current entry in entry field
        if user_text and self.save_flag == False: #not a numbered word or custom to save
            self.outlist.append(re.sub(self.save_key, user_text, self.current_word)) #substitute and append
        elif user_text and self.save_flag == True: #numbered word to save
            self.outlist.append(re.sub(self.save_key, user_text, self.current_word)) #substitute and append
            numword_dic[self.save_key]=user_text #save numbered word
            self.save_flag=False #turn off save flag
        else:
            pass
        self.next_prompt() #get next prompt

    def normalize_quotes(self, text): #normalizes curly quotes from Word docs when printing text files
        return text.replace('“', '"').replace('”', '"').replace("‘", "'").replace("’", "'")

    def third_window(self):  #decide whether to save filled output
        for widget in self.root.winfo_children():
            widget.destroy()
        display = scrolledtext.ScrolledText(self.root, width=80, height=20, font=("Arial", 12), bg="#9cc9e0", fg="black")
        display.pack(pady=20)
        self.filled = re.sub(r'\s([.,!?;:])', r'\1', ' '.join(self.outlist)) #rejoins output array without spacing between words and punctuation

        display.insert(tk.END, "\nHere is your filled Madlib:")
        if self.title: #display title if it exists
            display.insert(tk.END, "\n\n"+self.title)
        display.insert(tk.END, "\n\n"+self.filled)
        w = tk.Label(self.root, text='What would you like to do with it?', width=40, height=5, bg="#d0e7ff",
                     fg="black")
        w.pack(pady=5)
        # buttons for filled madlib save selection
        button_frame = tk.Frame(self.root)  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        # Plain text
        btn = tk.Button(button_frame, command=lambda: self.output_file_name(0), text="Save plain text", bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=0, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        #Word doc
        btn = tk.Button(button_frame, command=lambda: self.output_file_name(2), text="Save Word Doc", bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=1, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # Back to menu
        btn = tk.Button(button_frame, command=lambda: self.reset(), text="Return to menu",
                        bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
    def output_file_name(self,md): #output file namer
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        w = tk.Label(self.root, text='Enter the filename you\'d like to save to (no extension)',font=("Arial", 12, "bold"), width=80, height=10,
                     bg="#d0e7ff",
                     fg="black")
        w.pack(pady=10)
        # buttons for file naming menu selection
        button_frame = tk.Frame(self.root)  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        # buttons for file naming menu selection
        button_frame = tk.Frame(self.root)  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames


        self.input_entry = tk.Entry(self.root, font=("Arial", 14), width=80, bg="#d0e7ff", fg="black")
        self.input_entry.bind("<Return>",
                                  lambda event: self.output_save(md))  # allows the "enter" key to submit the keyword
        self.input_entry.pack(pady=10)
        self.input_entry.focus_set()  # automatically puts the cursor into the entry field
        self.submit_btn = tk.Button(self.root, text="Submit", command=lambda: self.output_save(md), bg="#3b9dd3",
                                        fg="white")
        self.submit_btn.pack(pady=10)


    def plain_file_name(self):
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        w = tk.Label(self.root, text='Enter the filename you\'d like to save to (no extension)', width=80, height=10,
                     bg="#d0e7ff",
                     fg="black")
        w.pack(pady=10)

        button_frame = tk.Frame(self.root)  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        self.input_entry = tk.Entry(self.root, font=("Arial", 14), width=80, bg="#d0e7ff", fg="black")
        self.input_entry.bind("<Return>",
                              lambda event: self.output_save(0))  # allows the "enter" key to submit the keyword
        self.input_entry.pack(pady=10)
        self.input_entry.focus_set()  # automatically puts the cursor into the entry field
        self.submit_btn = tk.Button(self.root, text="Submit", command=lambda: self.output_save(0), bg="#3b9dd3", fg="white") #submit button
        self.submit_btn.pack(pady=10)

    def invalid_html_window(self):
        for widget in self.root.winfo_children():
            widget.destroy()

        self.display = scrolledtext.ScrolledText(self.root, width=80, height=10, font=("Arial", 12), bg="#9cc9e0", fg="black")
        self.display.pack(pady=10)
        # define and create entry field for user's entry for a word
        self.input_entry = tk.Entry(self.root, font=("Arial", 14), width=80, bg="#d0e7ff", fg="black")
        self.input_entry.bind("<Return>", lambda event: self.process_invalid_html())  # allows the "enter" key to submit the keyword
        self.input_entry.pack(pady=10)
        self.input_entry.focus_set()  # automatically puts the cursor into the entry field

        self.submit_btn = tk.Button(self.root, text="Submit", command=self.process_invalid_html, bg="#3b9dd3", fg="white")
        self.submit_btn.pack(pady=10)

    def  html_replace(self): #todo: decide between original (had length bugs) and "_C." if "_C" then refactor keyword replacement to match this
        try:
            while True:
                self.current_word = next(self.html_words)
                final = '[undefined]'
                if re.findall(unnumbered, self.current_word) and not re.findall(numbered,
                                                                                self.current_word) and not re.findall(
                        customreg, self.current_word):
                    regkey = re.findall(unnumbered, self.current_word)
                    realkey = ''.join(regkey)
                    self.save_key=realkey
                    if realkey in generic_words:
                        htword = htmlsample.replace('underscript', generic_words[realkey])
                        final = self.current_word.replace(realkey, htword, 1)
                    elif realkey in ignored_words:
                        final = self.current_word
                    else:
                        self.display.insert(tk.END, f"{realkey} is not a valid keyword, what is it?:\n")
                        return

                elif re.findall(customreg, self.current_word) and not re.findall(numcustreg, self.current_word):
                    regkey = re.findall(customreg, self.current_word)
                    realkey = ''.join(regkey)
                    if realkey in self.custom:
                        htword = htmlsample.replace('underscript', self.custom[realkey])
                        final = self.current_word.replace(realkey, htword, 1)
                    else:
                        self.display.insert(tk.END, f"{realkey} is not a valid keyword, what is it?:\n")
                        return

                elif re.findall(numcustreg, self.current_word):
                    regkey = re.findall(numcustreg, self.current_word)
                    realkey = ''.join(regkey)
                    base = ''.join(re.findall(customreg, self.current_word))
                    num = ''.join(re.findall(r'_(\d+)', self.current_word))
                    #num = ''.join(re.findall(r'\d+', self.current_word))
                    htword = htmlsample.replace('underscript', self.custom[base] + ' (' + num + ')')
                    final = self.current_word.replace(realkey, htword, 1)

                elif re.findall(numbered, self.current_word):
                    regkey = re.findall(numbered, self.current_word)
                    regkeyb = re.findall(unnumbered, self.current_word)
                    realkey = ''.join(regkey)
                    base = ''.join(regkeyb)
                    num = ''.join(re.findall(r'\d+', realkey))
                    self.save_key = realkey
                    if base in generic_words:
                        htword = htmlsample.replace('underscript', generic_words[base] + ' (' + num + ')')
                        final = self.current_word.replace(realkey, htword, 1)
                    else:
                        self.display.insert(tk.END, f"{realkey} is not a valid keyword, what is it?:\n")
                        return
                else:
                    final = self.current_word

                self.htlist.append(final)

        except StopIteration:
            self.html_post_process()

    def advance_to_html(self):
        #user_input = self.raw_in
        #self.userMadlib = re.findall(r'/\w+\d*_[0-9]+|/\w+\d*|[^\s\w]|[\w]+', user_input)
        self.userMadlib=self.raw_in.split(' ')
        self.html_words = iter(self.userMadlib)
        self.invalid_html_window()
        self. html_replace()

    def process_invalid_html(self):
        user_text = self.input_entry.get().strip()
        self.input_entry.delete(0, tk.END)
        user_ht = htmlsample.replace('underscript', user_text)
        self.htlist.append(user_ht)
        self. html_replace()
    def html_post_process(self):
        for widget in self.root.winfo_children():
            widget.destroy()
        self.html_out = re.sub(r'\s([.,!?;:])', r'\1', ' '.join(self.htlist))
        if self.title:
            self.html_out = htmlhead.replace('heading', self.title, 1) + self.html_out + ' </p></body></html>'
        else:
            self.html_out = htmlhead_notitle + self.html_out + ' </p></body></html>'
        self.html_file_choice()

    def html_file_choice(self):
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        w = tk.Label(self.root, text='Before we print your Madlib, would you like to save it?', width=80, height=10, bg="#d0e7ff",font=("Arial", 12, "bold"),
                     fg="black")
        w.pack(pady=10)
        # buttons for Welcome menu selection
        button_frame = tk.Frame(self.root)  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        #Yes button
        btn = tk.Button(button_frame, command=lambda: self.output_file_name(1), text="Save", bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=0, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # no button
        btn = tk.Button(button_frame, command=lambda: self.html_view(), text="Print without saving", bg="#3b9dd3",
                        fg="white")  # defines each button with frame,
        btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        self.root.mainloop()  # deploys the GUI screen till closed

    def html_file_name(self):
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        w = tk.Label(self.root, text='Enter the filename you\'d like to save to (no extension)', width=80, height=10,
                     bg="#d0e7ff",
                     fg="black")
        w.pack(pady=10)
        # buttons for file naming menu selection
        button_frame = tk.Frame(self.root)  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        self.input_entry = tk.Entry(self.root, font=("Arial", 14), width=80, bg="#d0e7ff", fg="black")
        self.input_entry.bind("<Return>",
                              lambda event: self.output_save(1))  # allows the "enter" key to submit the keyword
        self.input_entry.pack(pady=10)
        self.input_entry.focus_set()  # automatically puts the cursor into the entry field
        self.submit_btn = tk.Button(self.root, text="Submit", command=lambda: self.output_save(1), bg="#3b9dd3",
                                    fg="white")
        self.submit_btn.pack(pady=10)
    def output_save(self,md): #todo: rename to generic save name
        base = self.input_entry.get().strip()
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        if md==0: #plain save
            if self.title:
                self.filled = self.title + "\n\n" + self.filled
            self.file_write(self.normalize_quotes(self.filled), base, 'outputs', '.txt')  # todo: add title
            w = tk.Label(self.root, text='Your filled mandlib has been saved to: ' + str(
                base) + '.txt in your \"outputs\" folder.\nWe hope you liked it!',font=("Arial", 12, "bold"),
                         width=80, height=10, bg="#d0e7ff",
                         fg="black")
            w.pack(pady=10)
            self.submit_btn = tk.Button(self.root, text="Back to menu", command=lambda: self.reset(), bg="#3b9dd3",
                                        fg="white")
            self.submit_btn.pack(pady=10)
        elif md==1: #html save
            
            self.file_write(self.html_out, base, 'outputs',
                            '.html')
            w = tk.Label(self.root, text='Your mandlib has been saved to: ' + str(
                base) + '.html in your \"outputs\" folder.\nNow let\'s print it!',
                         width=80, height=10, bg="#d0e7ff",font=("Arial", 12, "bold"),
                         fg="black")
            w.pack(pady=10)
            self.submit_btn = tk.Button(self.root, text="Let's Go", command=lambda: self.html_view(), bg="#3b9dd3",
                                        fg="white") #todo: add button grid and "back to menu" for file confirmations, beware of "laready has slaves" errors
            self.submit_btn.pack(pady=10)
        elif md==2: #word document outputs
            # Full path to save the document

            full_path = os.path.join('outputs', base + '.docx')

            doc = Document()

            if self.title:
                # Add title
                title_paragraph = doc.add_paragraph()
                title_run = title_paragraph.add_run(self.title)
                title_run.bold = True
                title_run.font.size = Pt(24)
                title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Add a blank line
                doc.add_paragraph()

            # Add body text
            body_paragraph = doc.add_paragraph()
            body_run = body_paragraph.add_run(self.filled)
            body_run.font.size = Pt(12)

            # Save the document
            doc.save(full_path)
            w = tk.Label(self.root, text='Your filled mandlib has been saved to: ' + str(
                base) + '.docx in your \"outputs\" folder.\nWe hope you liked it!',font=("Arial", 12, "bold"),
                         width=80, height=10, bg="#d0e7ff",
                         fg="black")
            w.pack(pady=10)
            self.submit_btn = tk.Button(self.root, text="Back to menu", command=lambda: self.reset(), bg="#3b9dd3",
                                        fg="white")
            self.submit_btn.pack(pady=10)
        elif md==3: #save and play inputs plain text
            self.file_write('<t>'+self.title+'</t>\n'+self.raw_in + '\n' + '<C>' + str(self.custom), base, 'inputs', '.txt')  # todo: decide if this should be done in file_choice
            w = tk.Label(self.root, text='Your mandlib has been saved to: ' + str(
                base) + '.txt in your \"inputs\" folder.\nNow we can Play!',
                         width=80, height=10, bg="#d0e7ff",font=("Arial", 12, "bold"),
                         fg="black")
            w.pack(pady=10)
            self.submit_btn = tk.Button(self.root, text="Let's Go", command=lambda: self.advance_to_second(),
                                        bg="#3b9dd3", fg="white")
            self.submit_btn.pack(pady=10)
        elif md==4: #save and play Word docx inputs
            full_path = os.path.join('inputs', base + '.docx')

            doc = Document()


            # Add title
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run('<t>'+self.title+'</t>')
            title_run.bold = True
            title_run.font.size = Pt(24)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add a blank line
            doc.add_paragraph()

            # Add body text
            body_paragraph = doc.add_paragraph()
            body_run = body_paragraph.add_run(self.raw_in)
            body_run.font.size = Pt(12)

            custom_paragraph = doc.add_paragraph()
            custom_run = custom_paragraph.add_run('<C>'+str(self.custom))
            custom_run.font.size= Pt(12)

            # Save the document
            doc.save(full_path)
            w = tk.Label(self.root, text='Your filled mandlib has been saved to: ' + str(
                base) + '.docx in your \"inputs\" folder.\nNow let\'s play it!',
                         width=80, height=10, bg="#d0e7ff",font=("Arial", 12, "bold"),
                         fg="black")
            w.pack(pady=10)
            self.submit_btn = tk.Button(self.root, text="Let's Go", command=lambda: self.advance_to_second(),
                                        bg="#3b9dd3", fg="white")
            self.submit_btn.pack(pady=10)
        else: #all invalids

            self.file_write(self.html_out + self.normalize_quotes(self.filled), base, 'outputs',
                            '.txt')  # todo: include selected (or loaded) title and formatting from terminal version
            w = tk.Label(self.root, text='Invalid save case found, please contact the developer.\n In the meantime, your mandlib has been saved to: ' + str(
                base) + '.txt in your \"outputs\" folder, but it won\'t be pretty.',
                         width=80, height=10, bg="#d0e7ff",font=("Arial", 12, "bold"),
                         fg="black")
            w.pack(pady=10)
            self.submit_btn = tk.Button(self.root, text="Ok", command=lambda: self.reset(), bg="#3b9dd3",
                                        fg="white")
            self.submit_btn.pack(pady=10)
    def html_view(self):

        self.root.destroy()  # closes the gui entirely todo: decide whether to quit the window here
        messagebox.showinfo("Thank you.", "We'll uploaded your madlib to your browser, you can print it from there.")
        with tempfile.NamedTemporaryFile("w", delete=False, suffix=".html") as f:
            f.write(self.html_out)
            webbrowser.open(f.name)
if __name__ == "__main__":
    root = tk.Tk()
    root.title("Madlib App")
    app = MadlibApp(root)
    root.mainloop()
