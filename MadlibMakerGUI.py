import tkinter as tk #main GUI tools
from tkinter import scrolledtext, messagebox, filedialog
from tkinter import font as tkfont
import webbrowser #for opening html files in browser
import tempfile #for WORD input file analysis
import re #regular expression library for substituting words
import os #for file and folder writing and reading

from long_strings_gui import * #collection of long strings
from docx import Document #for writing and reading word docs
from docx.shared import Pt #for word doc formatting
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT #for word doc formatting

numword_dic = {}
unnumbered = "(/...)"
numbered = "(/...[0-9]+)"
customreg = "(/ct[0-9]+)"
numcustreg = "(/ct[0-9]+_[0-9]+)"
tail = "(_[0-9])"
WarnRed="#F23F3F"

past_yellow= "#f5e97c"

light_yellow="#ffffd4"

dark_purp="#2d092d"

mid_purp="#676ec3"

light_purp="#aa8ddb"

fade_white="#948c95"

light_blue="#9bc7f5"

dark_red="#9e316e"

class MadlibApp:
    def __init__(self, root):
        self.root = root
        root.configure(bg=light_blue)  # pastel chrome/blue background
        self.reset()
#PREREQUISITES
    def reset(self):
        self.custom = {} #full custom dictionary
        self.custom_keys = [] #storage for recording custom keys detected
        self.custom_index = 0 #for custom detection
        self.save_flag = False #flag for saving numbered keywords
        self.folders() #creates input and output folders if not here
        self.outlist = [] #output array for filled madlibs
        self.htlist = [] #output array for html madlibs
        self.title = '' #madlib title
        self.load_mode = False  # decides mode (False=write or True=load)
        self.welcomeMenuHandler() #welcome menu

    def folders(self): #creates input and output folders if none exist
        if not os.path.isdir(os.path.join(os.getcwd(), "inputs")):
            os.mkdir(os.path.join(os.getcwd(), "inputs"))
        if not os.path.isdir(os.path.join(os.getcwd(), "outputs")):
            os.mkdir(os.path.join(os.getcwd(), "outputs"))

    def shadow_shine(self,frame,h,w):
        canvas = tk.Canvas(frame, width=w + 10, height=h + 10,
                           bg="#9bc7f5", highlightthickness=0)

        # Draw fake chrome outline
        canvas.create_rectangle(
            5, 5,
            5 + w,
            5 + h,
            fill="#9ac7f5",
            outline="#ffffff",
            width=2
        )
        canvas.create_line(5, 5 + h, 5 + w, 5 + h,
                           fill="#3a1c5d", width=3)  # bottom shadow
        canvas.create_line(5 + w, 5, 5 + w, 5 + h,
                           fill="#3a1c5d", width=3)  # right shadow
        canvas.create_line(5, 2 + h, 3 + w, 2 + h,
                           fill=darker("#9ac7f5"), width=3)  # bottom half-shadow
        canvas.create_line(2 + w, (h+5)*.5, 2 + w, h+2,
                           fill=darker("#9ac7f5"), width=3)  # right half-shadow
        return canvas


#todo: see if return functions can be moved out of class and onto another file, else, move pack commands to functions
    def hypno_title(self, text):
        w = tk.Label(
            self.root,
            text=text,
            font=("Courier New", 40, "bold"),
            fg=dark_red,
            bg="#9bc7f5",
            # height=5,
            width=20,
            padx=2,
            # pady=4
        )
        w.pack(padx=(10))

    def hypno_button(self, frame, text, command=None,color: str="#aa8ddb"): #todo: fully comment
        # Define font
        btn_font = ("Courier New", 10, "bold")

        # Measure the width of the text using a font object
        f = tkfont.Font(family="Courier New", size=10, weight="bold")
        text_width = f.measure(text)

        # Define padding and sizing
        padding_x = 20  # extra width around text
        padding_y = 12  # extra height
        button_width = text_width + padding_x
        button_height = 26 + padding_y

        canvas=self.shadow_shine(frame,button_height,button_width)
        # Add button
        btn = tk.Button(
            canvas,
            text=text,
            command=command,
            font=btn_font,
            bg=color,
            fg="white",
            activebackground="#9c6ddb",
            bd=0,
            relief="flat",
            highlightthickness=0
        )
        btn.place(x=8, y=8, width=button_width - 6, height=button_height - 8)

        return canvas

    def hypno_entry(self, w):
        entry = tk.Entry(
            self.root,
            font=("Courier New", 10),
            width=w,
            bg="#ffffd4",
            fg="black",
            bd=2)
        return entry

    def hypno_scroll(self, h, w, bg_color: str=light_yellow,fg_color=None):
        text_area = scrolledtext.ScrolledText(
            self.root,
            width=w,
            height=h,
            wrap="word",
            font=("Courier New", 10),
            bg=bg_color,
            fg=fg_color,
            padx=15
        )
        return text_area

    # Title style block (hypnospace-themed heading)
    def hypno_label(self, text, label_height, label_width, font_size, bg_color=past_yellow, fg_color=dark_red):
        frame = root
        # ------------------------------------------------------------------
        #  Font setup
        # ------------------------------------------------------------------
        btn_font = ("Courier New", font_size, "bold")
        fnt = tkfont.Font(family="Courier New",
                          size=font_size,
                          weight="bold")
        lines = text.splitlines() or [""]
        longest_line = max(lines, key=fnt.measure)
        line_width_px = fnt.measure(longest_line)
        line_height = fnt.metrics("linespace")
        n_lines = len(lines)
        #temporary nullify label height and width for new release todo: decide whether to remove inputs entirely
        label_width= None
        label_height = None
        # If caller passed None for size, fall back to auto‑measure
        if label_width is None:
            label_width  = line_width_px + 20   # horizontal padding
        if label_height is None:
            label_height = n_lines * line_height + 12 + 26  # vertical padding

        # ------------------------------------------------------------------
        #  Canvas with chrome outline
        # ------------------------------------------------------------------
        c = tk.Canvas(frame,
                      width=label_width + 10,
                      height=label_height + 10,
                      bg=frame["bg"],  # let the parent background show
                      highlightthickness=0)

        # light top/left todo: add inputs for additional hues
        top_hue=light_purp #hue of the top and left lines
        c.create_rectangle(5, 5,
                           5 + label_width,
                           5 + label_height,
                           fill=bg_color,
                           outline=top_hue, width=5)
        # deep shadow bottom/right
        c.create_line(7, 5 + label_height,
                      7 + label_width, 5 + label_height,
                      fill=darker(dark_red), width=5)
        c.create_line(5 + label_width, 7,
                      5 + label_width, 7 + label_height,
                      fill=darker(dark_red), width=5)

        # optional half‑shadow for extra gloss

        # ------------------------------------------------------------------
        #  Drop the real Label on top
        # ------------------------------------------------------------------
        lbl = tk.Label(c,
                       text=text,
                       font=btn_font,
                       bg=bg_color,
                       fg=fg_color,
                       bd=0,
                       highlightthickness=0)
        lbl.place(x=8, y=8,
                  width=label_width - 6,
                  height=label_height - 8)

        return c

    def hypno_header(self, text):
        self.header = tk.Frame(root, bg="#676ec3", height=30)
        self.header.pack(fill="x", side="top")

        header_label = tk.Label(
            self.header,
            text=text,
            bg="#676ec3",
            fg="white",
            font=("Courier New", 12, "bold")
        )

        header_label.pack(side="left", padx=10)


    def head_btn(self, text, command=None):
        exit_button = tk.Button(
            self.header,
            text=text,
            bg="#d95fbc",
            fg="white",
            font=("Courier New", 10, "bold"),
            relief="flat",
            command=command

        )
        exit_button.pack(side="right", padx=5)

    def dummyscreen(self,name): #placeholder for menus-to-be
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        text = "This is the screen for " + str(name)
        w = tk.Label(self.root, text=text, width=80, height=10, bg="#d0e7ff", fg="black")
        w.pack(pady=10)
        self.root.mainloop()  # deploys the GUI screen till closed

#MAIN MENU
    def welcomeMenuHandler(self): #Welcome menu
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets

        # welcome text
        self.hypno_header("Welcome to...")
        self.head_btn("Reset",command=lambda: self.reset())

        self.hypno_title("The Madlib\nMaker!")

        w = self.hypno_label('What would you like to do today?', None, None, 12)
        w.pack(pady=10)

        # buttons for Welcome menu selection
        button_frame = tk.Frame(self.root, bg="#9bc7f5")  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        # manual input button
        btn= self.hypno_button(button_frame, "Manual Input", command=lambda: self.setup_manual_window())
        btn.grid(row=1, column=0, padx=2, pady=2, sticky="ew")  # defines the button's location on the grid ("ew" centers all buttons to their grid position)
        # Load inputs button
        btn = self.hypno_button(button_frame, "Load a Madlib", command=lambda: self.load_input_file())
        btn.grid(row=1, column=2, padx=2, pady=2, sticky="ew")
        # Instruction menu button
        btn= self.hypno_button(button_frame, "Instructions", command=lambda: self.instruct_main())
        btn.grid(row=1, column=3, padx=2, pady=2,
                 sticky="ew")

        btn=self.hypno_button(self.root,"Credits & Special Thanks", command=lambda: self.credits())
        btn.pack()
        self.root.mainloop()  # deploys the GUI screen till closed

#INSTRUCTIONS
    def instruct_main(self): #Instructions top menu
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        self.hypno_header("Instructions")
        self.head_btn("Reset", command=lambda: self.reset())
        # Instructions text
        w = self.hypno_label('What would you like to learn about?', 2,40,12)
        w.pack(pady=10,padx=5)
        # buttons for instructions
        button_frame = tk.Frame(self.root, bg="#9bc7f5")  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        # How to play instruct
        btn= self.hypno_button(button_frame, 'How to play MadLibs',command=lambda: self.end_instructions(how_to_play_madlibs,False))
        btn.grid(row=1, column=0, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # How to write instruct button
        btn = self.hypno_button(button_frame, 'How to write madlibs',command=lambda: self.write_instructions_menu())
        btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # Loading madlib instruct button
        btn = self.hypno_button(button_frame, 'Loading a Madlib',command=lambda: self.end_instructions(loading_a_madlib,False))
        btn.grid(row=1, column=3, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # printing madlib instruct button
        btn = self.hypno_button(button_frame, 'Printing madlibs', command=lambda: self.end_instructions(printing_madlibs,False))
        btn.grid(row=1, column=4, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        #back to menu from instruct 1
        btn= self.hypno_button(button_frame, '< Back to main menu', command=lambda: self.reset())
        btn.grid(row=2,padx=2,pady=2,sticky="ew")
        self.root.mainloop()  # deploys the GUI screen till closed

    def write_instructions_menu(self):
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        self.hypno_header("Instructions")
        self.head_btn("Reset", command=lambda: self.reset())
        self.display= self.hypno_scroll(15,120)
        self.display.pack(pady=10, padx=5) #display widget spacing, horizontal and vertical
        self.display.insert(tk.END, how_to_write_a_madlib) #inserts "write madlibs" basic tutorial
        self.display.config(state='disabled')  # Make it read-only
        # buttons frame
        button_frame = tk.Frame(self.root, bg="#9bc7f5")   # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        # Generic words
        btn = self.hypno_button(button_frame, 'Generic words',
                                command=lambda: self.end_instructions(generic_words_list,True))
        btn.grid(row=1, column=0, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid

        #numbered words
        btn = self.hypno_button(button_frame, 'Numbered Words',
                                command=lambda: self.end_instructions(numbered_words,True))
        btn.grid(row=1, column=1, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # Custom words
        btn = self.hypno_button(button_frame, 'Custom words basics',
                                command=lambda: self.end_instructions(custom_words_basics,True))
        btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # numbered custom words
        btn = self.hypno_button(button_frame, 'Numbered Custom Words',
                                command=lambda: self.end_instructions(numbered_custom_words,True))
        btn.grid(row=1, column=3, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # Custom dic
        btn = self.hypno_button(button_frame, 'Preconfiguring custom words',
                                command=lambda: self.end_instructions(prewriting_custom_configurations,True))
        btn.grid(row=1, column=4, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # back button
        button_frame = tk.Frame(self.root, bg="#9bc7f5")  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        btn = self.hypno_button(button_frame, '< Back to Instructions',
                                command=lambda: self.instruct_main())
        btn.grid(row=2, column=1, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        #back to menu button
        btn = self.hypno_button(button_frame, '<< Back to main menu', command=lambda: self.reset())
        btn.grid(row=2, column=3, padx=2, pady=2, sticky="ew")  # defines the button's location on the grid

    def end_instructions(self,content,lv):#generic window for instructions at the end of the instruction tree.
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        self.hypno_header("Instructions")
        self.head_btn("Reset", command=lambda: self.reset())
        self.display=self.hypno_scroll(10, 80)
        self.display.pack(pady=10, padx=5)
        self.display.insert(tk.END, content)
        self.display.config(state='disabled')  # Make it read-only

        button_frame = tk.Frame(self.root, bg="#9bc7f5")  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames

        if lv: #second level of menus, currently only used in specifics of writing madlibs, change to an int if more are added
            # back to write instructions button
            btn = self.hypno_button(button_frame, '< Back to writing Madlibs',command=lambda: self.write_instructions_menu())
            btn.grid(row=1, column=1, padx=2, pady=2, sticky="ew")
            #back to instructions from lv2 end instructions
            btn = self.hypno_button(button_frame, '<< Back to Instructions',  command=lambda: self.instruct_main())
            btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
            #back to menu from lv2 end instructions
            btn = self.hypno_button(button_frame, '<<< Back to main menu', command=lambda: self.reset())
            btn.grid(row=1, column=3, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        else: #normal end instructions, only back to menu and back to instructions
            btn = self.hypno_button(button_frame, '< Back to Instructions', command=lambda: self.instruct_main())
            btn.grid(row=1, column=1, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
            btn = self.hypno_button(button_frame, '<< Back to main menu', command=lambda: self.reset())
            btn.grid(row=1, column=3, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
    def credits(self): #place for credits and special thanks
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        self.hypno_header("Credits")
        self.head_btn("Reset", command=lambda: self.reset())
        self.display=self.hypno_scroll(10, 80)
        self.display.pack(pady=10, padx=5)
        self.display.insert(tk.END, credits)
        self.display.config(state='disabled')  # Make it read-only
        btn = self.hypno_button(self.root, '<< Back to main menu', command=lambda: self.reset())
        btn.pack()
#READING FILES
    def load_input_file(self):
        # Clear existing widgets if necessary
        for widget in self.root.winfo_children():
            widget.destroy()
        self.hypno_header("Load Input File")
        self.head_btn("Reset", command=lambda: self.reset())
        self.load_mode = True #begin load mode
        label = self.hypno_label("Select a file to load:",None,None,20)
        label.pack(pady=10)

        inputs_path = os.path.join(os.getcwd(), "inputs")

        if not os.path.exists(inputs_path): #double check for inputs folder todo: add this check for every file write operation
            os.makedirs(inputs_path)

        files = [f for f in os.listdir(inputs_path) if f.endswith('.txt') or f.endswith('.docx')] #grabs all files in inputs

        btn = self.hypno_button(self.root,"Browse for File",command=self.browse_and_load_file)
        btn.pack()
        button_frame = tk.Frame(self.root, bg="#9bc7f5")  # defines the button frame
        button_frame.pack()


        row = 0  # initializes row and column counters for button grid
        col = 0
        for filename in files: #generate buttons for all files
            full_path = os.path.join(inputs_path, filename)
            btn = self.hypno_button(button_frame, filename, command=lambda path=full_path: self.process_input_file(path))
            btn.grid(row=row, column=col, sticky="ew")  # defines the button's location on the grid ("ew" centers all buttons to their grid position)
            col += 1
            if col >= 5:  # max number of columns, then new row
                button_frame = tk.Frame(self.root, bg="#9bc7f5")  # New frame to ensure center justification
                button_frame.pack()
                col = 0
                row += 1
        if not files: #edge case of no valid input files
            no_files_label = self.hypno_label("No input files found in the 'inputs' folder.\nPlace one into the folder and click \"Refresh\" or Browse for it instead.",None,None,12)
            no_files_label.pack(pady=10)
        else:
            no_files_label = self.hypno_label(
                "Don't see your file?\nPlace it in the \"inputs\" folder and click \"Refresh\" or Browse for it instead", None,
                None, 12)
            no_files_label.pack(pady=5)
        rfbtn = self.hypno_button(self.root, "refresh", command=lambda: self.load_input_file())
        rfbtn.pack(pady=10)

    def browse_and_load_file(self):
        # Open a file dialog for the user to select a file
        path = filedialog.askopenfilename(
            title="Select a Madlib File",
            filetypes=[("All Files", "*.*"),("Text Files", "*.txt"), ("Word Documents", "*.docx")]
        )

        # If a file was selected, process it
        if path:
            self.process_input_file(path)

    def process_input_file(self, path): #todo: comment lines
        custom_dict = {} #temporary custom dic storage
        title_text = "" #temporary title storage
        file_path = path
        try:
            # Read file content
            if file_path.endswith(".txt"):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_path.endswith(".docx"):
                doc = Document(file_path)
                content = "\n".join([p.text for p in doc.paragraphs])
            else:
                messagebox.showerror("Unsupported Format", "Only .txt and .docx files are supported.")
                return

            # Extract title if present
            if "<t>" in content and "</t>" in content:
                start = content.index("<t>") + len("<t>")
                end = content.index("</t>")
                title_text = content[start:end].strip()

                # Remove title from content
                content = content[:content.index("<t>")] + content[end + len("</t>"):]

            # Extract custom dictionary
            if "<C>" in content:
                c_index = content.index("<C>")
                madlib_text = content[:c_index].strip()
                dict_text = content[c_index + len("<C>"):].strip()
                try:
                    custom_dict = eval(dict_text)
                except Exception as e:
                    messagebox.showerror("Custom Dict Error", f"Failed to parse custom dictionary:\n{e}")
                    return
            else:
                madlib_text = content.strip()

            self.title = title_text
            self.custom = custom_dict
            self.raw_in=madlib_text
            self.advance_from_manual()

        except Exception as e:
            messagebox.showerror("File Error", f"Could not read file:\n{e}")

#MANUAL INPUT
    def setup_manual_window(self): #setup manual input window
        for widget in self.root.winfo_children():
            widget.destroy()
        self.hypno_header("Manual Input")
        self.head_btn("Reset", command=lambda: self.reset())
        l = self.hypno_label("Type your Madlib below",None,None,20)
        l.pack(pady=5)

        self.input_entry = self.hypno_entry(80)
        self.input_entry.pack(pady=10)# sets the verticle spacing given between the input field and other successive elements
        self.input_entry.focus_set()  # This sets focus so the cursor appears in the field

        l = self.hypno_label("Use these buttons to drop a keyword",None,None,12)
        l.pack(pady=5)
        button_frame = tk.Frame(self.root, bg="#9bc7f5")  # defines the button frame
        button_frame.pack(pady=5, padx=15)

        row = 0 #initializes row and column counters for button grid
        col = 0
        for key, label in generic_words.items(): #for every generic word. Grabs the key (/adj) and label (adjective)
            btn = tk.Button(button_frame, text=label, command=lambda k=key: self.insert_keyword(k), bg="#aa8ddb", fg="white") #defines each button with frame todo: Consider a unique color for this
            btn.grid(row=row, column=col, padx=2, pady=2, sticky="ew") #defines the button's location on the grid
            col += 1
            if col >= 10: #ten columns, then new row
                col = 0
                row += 1

        #define submit button, command to advance to next screen
        self.submit_btn = self.hypno_button(self.root,"Submit",command=self.advance_from_manual)
        self.submit_btn.pack(pady=10)
        #define display
        self.display = self.hypno_scroll(10,80,dark_purp,"#948c95")
        self.display.pack(pady=10)
        self.display.configure(state='disabled')#disables edits
        self.input_entry.bind("<KeyRelease>", self.sync_display) #syncs display on every key release

    def sync_display(self, event=None): #syncs text between manual entry and display todo: verify why "event" is needed but not used in function
        self.display.configure(state='normal') #re-enables edits for syncing
        self.display.delete(1.0, tk.END) #clears the display before updating with current text (from row 1 char 0 to the end)
        self.display.insert(tk.END, self.input_entry.get()) #inserts text from the end (which is the start of the field after delete) to
        self.display.configure(state='disabled')  # re-disables after syncing 

    def insert_keyword(self, keyword):#inserts a keyword into active input field
        self.input_entry.insert(tk.END, keyword) #inputs keyword without a space (do not have both lines active)
        self.sync_display() #adds the keyword to the display

    def advance_from_manual(self): #handles input data
        # grabs manual input if manual
        if self.load_mode == False:#manual input
            self.raw_in = self.input_entry.get()
        elif self.load_mode == True: #from file
            pass
        else:
            self.dummyscreen('Invalid mode for advance_from_manual()')

        self.userMadlib= self.raw_in.split(' ') #splits madlib body into array by spaces
        self.prompt_words = iter(self.userMadlib) #sets up iteration

        # Extract and sort custom keys
        all_custom_matches = re.findall(r'/ct(\d+)(?:_\d+)?', self.raw_in)
        self.custom_keys = sorted(set(f"/ct{id}" for id in all_custom_matches), key=lambda x: int(x[3:]))

        if self.custom_keys and self.custom=={}:#if customs detected
            self.custom_configure_window() #configure customs, file_choice will be executed as well
        else: #if no others are needed
            self.title_check()

#TITLE MAKERS
    def title_check(self): #check if title is set
        if self.title:
            self.file_choice() #advance to input preview
        else:
            self.title_write() #write title

    def title_write(self): #write title
        for widget in self.root.winfo_children():
            widget.destroy()
        self.hypno_header("Title your Madlib")
        self.head_btn("Reset", command=lambda: self.reset())
        self.hypno_label("Would you like to title your Madlib?",None,None,12).pack()
        self.input_entry = self.hypno_entry(80)
        self.input_entry.bind("<Return>",
                              lambda event: self.title_saver()) #"enter" key submits title
        self.input_entry.pack(pady=10)
        self.input_entry.focus_set()  # automatically puts the cursor into the entry field
        button_frame = tk.Frame(self.root, bg="#9bc7f5")  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        # Submit button to save title
        btn = self.hypno_button(button_frame,"Submit",command=lambda: self.title_saver())
        btn.grid(row=1, column=0, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # Skip button, go to input preview
        btn = self.hypno_button(button_frame, "Skip", command=lambda: self.file_choice())
        btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")

    def title_saver(self): #saves title
        self.title = self.input_entry.get().strip()
        self.file_choice()

#CUSTOM PROCESSING
    def custom_configure_window(self): #custom word configure window
        for widget in self.root.winfo_children():
            widget.destroy()
        self.hypno_header("Custom configurator")
        self.head_btn("Reset", command=lambda: self.reset())
        self.hypno_label("Custom words detected, please configure them.", None, None, 12).pack()
        self.display = self.hypno_scroll(10,80)
        self.display.pack(pady=10)

        self.custom_entry = self.hypno_entry(40)
        self.custom_entry.bind("<Return>",
                               lambda event: self.save_custom_word())  # allows the "enter" key to submit the custom word
        self.custom_entry.pack(pady=10)
        self.custom_button = self.hypno_button(self.root,"Enter",command=self.save_custom_word)
        self.custom_button.pack(pady=5)

        self.prompt_next_custom() #prompt the next custom word in the display
        self.root.mainloop()  # deploys the GUI screen till closed

    def prompt_next_custom(self): #promts each custom word in display
        if self.custom_index < len(self.custom_keys):
            current_key = self.custom_keys[self.custom_index] #grab the custom key based on sequential index
            self.display.insert(tk.END, f"Custom {current_key[3:]}: ") #inserts each prompt with the entry based on ID number
        else:
            self.title_check() #with all customs configured, advance to title check

    def save_custom_word(self): #saves custom word to custom dictionary
        current_key = self.custom_keys[self.custom_index]
        user_input = self.custom_entry.get().strip()
        if user_input:
            self.custom[current_key] = user_input #adds user input to custom dictionary for current key
            self.custom_entry.delete(0, tk.END) #deletes user input from the input field
            self.custom_index += 1 #updates the custom keys index
            self.display.insert(tk.END, f"{user_input}\n") #append user input to display for reference
            self.prompt_next_custom() #prompt the next custom word to be configured

#UNFILLED WINDOW
    def file_choice(self):
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        self.hypno_header("Madlib Preview")
        self.head_btn("Reset", command=lambda: self.reset())
        w = self.hypno_label("What would you like to do?", None, None, 12)
        button_frame = tk.Frame(self.root, bg="#9bc7f5")  # defines the button frame
        self.hypno_label("Your Madlib is ready!",None, None, 20).pack(pady=5)



        if not self.load_mode: #manual input, show it, no hide needed

            self.display = self.hypno_scroll(20, 80)
            self.display.pack(pady=20, padx=5)

            w.pack(pady=5)
            button_frame.pack(pady=5)  # for all button frames
            self.display.insert(tk.END, self.title + "\n\n"+self.raw_in)

            btn = self.hypno_button(button_frame,"Save plain text input",command=lambda: self.output_file_name(3))
            btn.grid(row=1, column=0, padx=2, pady=2,
                     sticky="ew")  # defines the button's location on the grid

            btn = self.hypno_button(button_frame,"Save Word doc input",command=lambda: self.output_file_name(4))
            btn.grid(row=1, column=1, padx=2, pady=2,
                     sticky="ew")  # defines the button's location on the grid
        else: #load, use button to reveal
            self.hypno_label("If you are the author, click \"SPOIL\" to preview it.\nOtherwise, we recommend you click \"Play\" to play it blind.",None,None,12).pack(pady=5)
            # self.display = scrolledtext.ScrolledText(self.root, width=80, height=20, font=("Arial", 12), bg="#9cc9e0", fg="black", wrap=tk.WORD)
            self.display = self.hypno_scroll(None, 60)
            self.display.pack(pady=20)
            self.display.insert(tk.END,"???")

            w.pack(pady=5)
            button_frame = tk.Frame(self.root, bg="#9bc7f5")  # defines the button frame
            button_frame.pack(pady=5)  # for all button frames

            self.spbtn = self.hypno_button(button_frame,"SPOIL",command=lambda: self.spoiler(),color="#F23F3F")
            self.spbtn.grid(row=1, column=4, padx=2, pady=2,
                     sticky="ew")



        # play without saving button

        btn = self.hypno_button(button_frame,"Play", command=lambda: self.advance_to_play())
        btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid

        btn = self.hypno_button(button_frame,"Print Physical",command=lambda: self.advance_to_html())
        btn.grid(row=1, column=3, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        self.root.mainloop()  # deploys the GUI screen till closed

    def spoiler(self): #inserts madlib input text to display if the user chooses to spoil a loaded madlib
        self.load_mode=False
        self.file_choice()

#PLAY WINDOW
    def advance_to_play(self): #advance to second window, start word prompting when selection is made
        self.play_window()
        self.next_prompt()

    def play_window(self): #generates Madlib play/fill window
        for widget in self.root.winfo_children():widget.destroy()
        self.hypno_header("Let's play!")
        self.head_btn("Reset", command=lambda: self.reset())
        #tk.Label(self.root, text="Give us a/an:", font=("Arial", 12, "bold")).pack()
        self.hypno_label("Give us a/an:", None, None, 12).pack()

        self.display= self.hypno_scroll(10,80)
        self.display.pack(pady=10, padx=5)

        self.hypno_label("Your Entry",None,None,12).pack()

        self.input_entry = self.hypno_entry(40)
        self.input_entry.bind("<Return>",
                              lambda event: self.process_next_keyword())  # allows the "enter" key to submit the keyword
        self.input_entry.pack(pady=10)
        self.input_entry.focus_set()  # automatically puts the cursor into the entry field

        self.submit_btn=self.hypno_button(self.root,"Submit",command=self.process_next_keyword)
        self.submit_btn.pack(pady=10)

    def process_next_keyword(self):
        user_text = self.input_entry.get().strip() #grab user text
        self.display.insert(tk.END, f"{user_text}\n") #display next to prompt in display
        self.input_entry.delete(0, tk.END) #delete current entry in entry field
        if user_text and self.save_flag == False: #not a numbered word or custom to save
            self.outlist.append(re.sub(self.save_key, user_text, self.current_word)) #substitute and append
        elif user_text and self.save_flag == True: #numbered word to save
            self.outlist.append(re.sub(self.save_key, user_text, self.current_word)) #substitute and append
            numword_dic[self.save_key]=user_text #save numbered word
            self.save_flag=False #turn off save flag
        else:
            pass
        self.next_prompt() #get next prompt

    def next_prompt(self): #iterates through each word of the madlib array, returns when user input is needed todo: complete comments
        try:
            self.current_word = next(self.prompt_words)

            #Generic word unnumbered
            if re.findall(unnumbered, self.current_word) and not re.findall(numbered,self.current_word) and not re.findall(customreg, self.current_word):
                regkey = re.findall(unnumbered, self.current_word)
                realkey = ''.join(regkey)
                self.save_key = realkey
                if realkey in generic_words:
                    self.display.insert(tk.END, f"{generic_words[realkey]}: ")
                    return
                elif realkey in ignored_words:
                    return
                else:
                    self.display.insert(tk.END,
                                        f"{realkey} is not a valid keyword, what would you like to replace it with?: ")
                    return

            elif re.findall(customreg, self.current_word) and not re.findall(numcustreg, self.current_word):
                # custom
                regkey = re.findall(customreg, self.current_word)
                realkey = ''.join(regkey)
                self.save_key = realkey
                if realkey in self.custom:
                    self.display.insert(tk.END, f"{self.custom[realkey]}: ")
                    return
                elif realkey not in self.custom:
                    self.display.insert(tk.END,f"{realkey} hasn't been configured, what would you like to replace it with?: ")
                    return

            elif re.findall(numcustreg, self.current_word):
                # numbered custom

                regkey = re.findall(numcustreg, self.current_word)
                realkey = ''.join(regkey)

                base = re.findall(customreg, self.current_word)
                base = ''.join(base)
                self.save_key = realkey
                if realkey not in numword_dic:# todo: fix unconfigged numbered customs, try adding "and base in self.custom"
                    # numbered cust unsaved
                    base = re.findall(customreg, self.current_word)
                    base = ''.join(base)
                    self.display.insert(tk.END, f"{self.custom[base]} : ")
                    self.save_flag = True
                    return
                elif realkey in numword_dic:
                    self.current_word = re.sub(realkey, numword_dic[realkey], self.current_word)
                    # numbered cust saved
                elif base not in self.custom:
                    # numbered cust unsaved
                    self.display.insert(tk.END,
                                        f"{realkey} hasn't been configured, what would you like to replace it with?: ")
                    self.save_flag = True
                    return
                else:
                    # others
                    self.display.insert(tk.END,
                                        f"{realkey} is not a valid keyword, what would you like to replace it with?: ")
                    return

            elif re.findall(numbered, self.current_word):
                # numbered
                regkey = re.findall(numbered, self.current_word)
                regkeyb = re.findall(unnumbered, self.current_word)
                realkey = ''.join(regkey)
                base = ''.join(regkeyb)
                self.save_key = realkey
                if realkey not in numword_dic.keys() and base in generic_words.keys():
                    self.display.insert(tk.END, f"{generic_words[base]}: ")
                    self.save_flag = True
                    return


                elif realkey not in numword_dic.keys() and base not in generic_words.keys():
                    # numbered unsaved

                    self.display.insert(tk.END,f"{realkey} is not a valid keyword, what would you like to replace it with?: ")
                    return
                elif realkey in numword_dic.keys():
                    self.current_word =  re.sub(realkey, numword_dic[realkey], self.current_word)
            else:
                pass
            self.outlist.append(self.current_word)
            self.next_prompt()
        except StopIteration:
            self.filled_window()

#FILLED WINDOW
    def filled_window(self):  #decide whether to save filled output
        for widget in self.root.winfo_children():
            widget.destroy()
        self.hypno_header("Madlib Complete")
        self.head_btn("Reset", command=lambda: self.reset())
        self.hypno_label("Your Madlib is complete!", None, None, 20).pack(pady=5)
        #display = scrolledtext.ScrolledText(self.root, width=80, height=20, font=("Arial", 12), bg="#9cc9e0", fg="black", wrap=tk.WORD)
        display = self.hypno_scroll(20,80)
        display.pack(pady=20)
        self.filled = re.sub(r'\s([.,!?;:])', r'\1', ' '.join(self.outlist)) #rejoins output array without spacing between words and punctuation

        # display.insert(tk.END, "\nHere is your filled Madlib:")
        if self.title: #display title if it exists
            display.insert(tk.END, "\n\n"+self.title)
        display.insert(tk.END, "\n\n"+self.filled)
        # w = tk.Label(self.root, text='What would you like to do with it?', width=40, height=5, bg="#d0e7ff",
        #              fg="black")
        w = self.hypno_label("What would you like to do with it?",None,None,12)
        w.pack(pady=5)
        # buttons for filled madlib save selection
        button_frame = tk.Frame(self.root, bg="#9bc7f5")  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        # Plain text

        btn = self.hypno_button(button_frame, "Save plain text", command=lambda: self.output_file_name(0))
        btn.grid(row=1, column=0, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        #Word doc
        btn = self.hypno_button(button_frame, "Save Word Doc", command=lambda: self.output_file_name(2))
        btn.grid(row=1, column=1, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # Back to menu
        # btn = tk.Button(button_frame, command=lambda: self.reset(), text="Return to menu",
        #                 bg="#3b9dd3",
        #                 fg="white")  # defines each button with frame,
        btn = self.hypno_button(button_frame, "Return to menu", command=lambda: self.reset())
        btn.grid(row=1, column=2, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid

    def output_file_name(self,md): #output file namer
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        self.hypno_header("Save file name")
        self.head_btn("Reset", command=lambda: self.reset())

        w = self.hypno_label("Enter the filename you\'d like to save to (no extension)",None,None,12)
        w.pack(pady=10)

        self.input_entry = self.hypno_entry(80)
        self.input_entry.bind("<Return>",
                                  lambda event: self.save_file(md))  # allows the "enter" key to submit the keyword
        self.input_entry.pack(pady=10)
        self.input_entry.focus_set()  # automatically puts the cursor into the entry field

        self.submit_btn = self.hypno_button(self.root,"Submit",command=lambda: self.save_file(md))
        self.submit_btn.pack(pady=10)

#SAVING OUTPUTS
    def save_file(self,md): # saving input and output
        base = self.input_entry.get().strip()
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        self.hypno_header("File Saved")
        self.head_btn("Reset", command=lambda: self.reset())
        if md==0: #plain save
            if self.title:
                self.filled = self.title + "\n\n" + self.filled
            self.file_write(self.normalize_quotes(self.filled), base, 'outputs', '.txt')

            w = self.hypno_label("Your filled madlib has been saved to:\n" + str(base) + ".txt\nin your \"outputs\" folder.\nWe hope you liked it!",None,None,12)
            w.pack(pady=10)

            self.hypno_button(self.root,"Back to Menu",command=lambda: self.reset()).pack(pady=10)
        elif md==1: #html save

            self.file_write(self.html_out, base, 'outputs','.html')

            w= self.hypno_label('Your madlib has been saved to: ' + str(base) + '.html in your \"outputs\" folder.\nNow let\'s print it!',None,None,12)
            w.pack(pady=10)

            self.submit_btn = self.hypno_button(self.root, "Let's Go",command=lambda: self.html_view())#todo: add button grid and "back to menu" for file confirmations, beware of "aready has slaves" errors
            self.submit_btn.pack(pady=10)
        elif md==2: #word document outputs
            # Full path to save the document

            full_path = os.path.join('outputs', base + '.docx')

            doc = Document()

            if self.title:
                # Add title
                title_paragraph = doc.add_paragraph()
                title_run = title_paragraph.add_run(self.title)
                title_run.bold = True
                title_run.font.size = Pt(24)
                title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                # Add a blank line
                doc.add_paragraph()

            # Add body text
            body_paragraph = doc.add_paragraph()
            body_run = body_paragraph.add_run(self.filled)
            body_run.font.size = Pt(12)

            # Save the document
            doc.save(full_path)

            w = self.hypno_label('Your filled madlib has been saved to:\n' + str(base) + '.docx\nin your \"outputs\" folder.\nWe hope you liked it!',None,None,12)
            w.pack(pady=10)
            self.hypno_button(self.root,"Back to Menu",command=lambda: self.reset()).pack(pady=10)
        elif md==3: #save and play inputs plain text
            self.file_write('<t>'+self.title+'</t>\n'+self.raw_in + '\n' + '<C>' + str(self.custom), base, 'inputs', '.txt')

            w = self.hypno_label('Your madlib has been saved to:\n' + str(base) + '.txt\nin your \"inputs\" folder.\nNow we can Play!',None,None,12) #todo: consider unique frames for filename and message
            w.pack(pady=10,padx=5)
            self.submit_btn = self.hypno_button(self.root,"Let's go!", command=lambda: self.advance_to_play()).pack(pady=10)
        elif md==4: #save and play Word docx inputs
            full_path = os.path.join('inputs', base + '.docx')

            doc = Document()


            # Add title
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run('<t>'+self.title+'</t>')
            title_run.bold = True
            title_run.font.size = Pt(24)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Add a blank line
            doc.add_paragraph()

            # Add body text
            body_paragraph = doc.add_paragraph()
            body_run = body_paragraph.add_run(self.raw_in)
            body_run.font.size = Pt(12)

            custom_paragraph = doc.add_paragraph()
            custom_run = custom_paragraph.add_run('<C>'+str(self.custom))
            custom_run.font.size= Pt(12)

            # Save the document
            doc.save(full_path)

            w = self.hypno_label('Your filled madlib has been saved to:\n' + str(
                base) + '.docx\nin your \"inputs\" folder.\nNow let\'s play it!',None,None,12)
            w.pack(pady=10)

            self.submit_btn = self.hypno_button(self.root,"Let's Go",command=lambda: self.advance_to_play())
            self.submit_btn.pack(pady=10)
        else: #all invalids

            self.file_write(self.html_out + self.normalize_quotes(self.filled), base, 'outputs',
                            '.txt')  #normalize quotes for text

            w = self.hypno_label('Invalid save case found, please contact the developer.\n In the meantime, your madlib has been saved to:\n'+ str(
                base)+'.txt\nin your \"outputs\" folder.',None,None,12)
            w.pack(pady=10)

            self.submit_btn = self.hypno_button(self.root,"Ok",command=lambda: self.reset())
            self.submit_btn.pack(pady=10)

    def normalize_quotes(self, text): #normalizes curly quotes from Word docs when printing text files todo: convert to an external function
        return text.replace('“', '"').replace('”', '"').replace("‘", "'").replace("’", "'")

    def file_write(self, content, name_of_file, path, ext): #writes the file todo: Consider making external
        completeName = os.path.join(path, name_of_file + ext) #filepath with name and extension
        f = open(completeName, "w") #open file with path
        f.write(content) #insert contents
        f.close() #close the file

#HTML PROCESSING/CONVERSION
    def advance_to_html(self): #prepare and send madlib input to convert to HTML
        self.userMadlib=self.raw_in.split(' ') #split madlib input by spaces
        self.html_words = iter(self.userMadlib) #set up iteration
        self.invalid_html_window() # prepair window for user inputs when keys invalid
        self. html_replace() # run html replacement

    def html_replace(self): #todo: fully comment
        try:
            while True:
                self.current_word = next(self.html_words)
                if re.findall(unnumbered, self.current_word) and not re.findall(numbered, self.current_word) and not re.findall(
                        customreg, self.current_word):
                    regkey = re.findall(unnumbered, self.current_word)
                    realkey = ''.join(regkey)
                    self.save_key=realkey
                    if realkey in generic_words:
                        htword = htmlsample.replace('underscript', generic_words[realkey])
                        final = self.current_word.replace(realkey, htword, 1)
                    elif realkey in ignored_words:
                        final = self.current_word
                    else:
                        self.display.insert(tk.END, f"{realkey} is not a valid keyword, what is it?:\n")
                        return

                elif re.findall(customreg, self.current_word) and not re.findall(numcustreg, self.current_word):
                    regkey = re.findall(customreg, self.current_word)
                    realkey = ''.join(regkey)
                    if realkey in self.custom:
                        htword = htmlsample.replace('underscript', self.custom[realkey])
                        final = self.current_word.replace(realkey, htword, 1)
                    else:
                        self.display.insert(tk.END, f"{realkey} is not a valid keyword, what is it?:\n")
                        return

                elif re.findall(numcustreg, self.current_word):
                    regkey = re.findall(numcustreg, self.current_word)
                    realkey = ''.join(regkey)
                    base = ''.join(re.findall(customreg, self.current_word))
                    num = ''.join(re.findall(r'_(\d+)', self.current_word))
                    htword = htmlsample.replace('underscript', self.custom[base] + ' (' + num + ')')
                    final = self.current_word.replace(realkey, htword, 1)

                elif re.findall(numbered, self.current_word):
                    regkey = re.findall(numbered, self.current_word)
                    regkeyb = re.findall(unnumbered, self.current_word)
                    realkey = ''.join(regkey)
                    base = ''.join(regkeyb)
                    num = ''.join(re.findall(r'\d+', realkey))
                    self.save_key = realkey
                    if base in generic_words:
                        htword = htmlsample.replace('underscript', generic_words[base] + ' (' + num + ')')
                        final = self.current_word.replace(realkey, htword, 1)
                    else:
                        self.display.insert(tk.END, f"{realkey} is not a valid keyword, what is it?:\n")
                        return
                else:
                    final = self.current_word

                self.htlist.append(final)

        except StopIteration:
            self.html_post_process()

    def invalid_html_window(self):
        for widget in self.root.winfo_children():
            widget.destroy()
        self.hypno_header("Invalid Keyword")
        self.head_btn("Reset", command=lambda: self.reset())

        self.display = self.hypno_scroll(10,80)
        self.display.pack(pady=10)
        # define and create entry field for user's entry for a word
        self.input_entry = self.hypno_entry(80)
        self.input_entry.bind("<Return>", lambda event: self.process_invalid_html())  # allows the "enter" key to submit the keyword
        self.input_entry.pack(pady=10)
        self.input_entry.focus_set()  # automatically puts the cursor into the entry field

        self.submit_btn = self.hypno_button(self.root,"Submit", command=self.process_invalid_html)
        self.submit_btn.pack(pady=10)

    def process_invalid_html(self): #grabs user input submitted when key is found invalid during html replacement
        user_text = self.input_entry.get().strip()
        self.input_entry.delete(0, tk.END)
        user_ht = htmlsample.replace('underscript', user_text) #substitute the word "underscript" in html sample for word under the underset
        self.htlist.append(user_ht)
        self. html_replace()

    def html_post_process(self):
        self.html_out = re.sub(r'\s([.,!?;:])', r'\1', ' '.join(self.htlist)) #join body with punctuation
        if self.title: #if title exists, place it in heading format string
            self.html_out = htmlhead.replace('heading', self.title, 1) + self.html_out + ' </p></body></html>'
        else:
            self.html_out = htmlhead_notitle + self.html_out + ' </p></body></html>'
        self.html_file_choice() #prompt save decision

    def html_file_choice(self): #screen for deciding whether to save html output
        for widget in self.root.winfo_children(): widget.destroy()  # removes pre-existing widgets
        self.hypno_header("Printing")
        self.head_btn("Reset", command=lambda: self.reset())
        # w = tk.Label(self.root, text='Before we print your Madlib, would you like to save it?', width=80, height=10, bg="#d0e7ff",font=("Arial", 12, "bold"), fg="black")
        w = self.hypno_label("Before we print your Madlib, would you like to save it?",None,None,12)
        w.pack(pady=10)
        # button frame
        button_frame = tk.Frame(self.root, bg="#9bc7f5")  # defines the button frame
        button_frame.pack(pady=5)  # for all button frames
        #Save button
        btn = self.hypno_button(button_frame,"Save",command=lambda: self.output_file_name(1))
        btn.grid(row=1, column=0, padx=2, pady=2,
                 sticky="ew")  # defines the button's location on the grid
        # Print without save button
        btn = self.hypno_button(button_frame, "Print without Saving",command=lambda: self.html_view())
        btn.grid(row=1, column=2, padx=2, pady=2, sticky="ew")  # defines the button's location on the grid
        self.root.mainloop()  # deploys the GUI screen till closed

    def html_view(self): #notification of opening HTML in browser
        self.root.destroy()  # closes the gui entirely
        messagebox.showinfo("Thank you.", "We'll uploaded your madlib to your browser, you can print it from there.")
        with tempfile.NamedTemporaryFile("w", delete=False, suffix=".html") as f:
            f.write(self.html_out)
            webbrowser.open(f.name)

def darker(hex_color: str, factor: float = 0.85) -> str:
    if not hex_color.startswith("#") or len(hex_color) != 7:
        raise ValueError("Colour must be in the form '#RRGGBB'.")

    hex_color = hex_color.lstrip("#")
    r, g, b = (int(hex_color[i : i + 2], 16) for i in (0, 2, 4))

    r = max(0, min(255, int(r * factor)))
    g = max(0, min(255, int(g * factor)))
    b = max(0, min(255, int(b * factor)))

    return f"#{r:02x}{g:02x}{b:02x}"

if __name__ == "__main__":
    root = tk.Tk() #initiate tkinter root
    root.title("Madlib App") #title root (all windows)
    app = MadlibApp(root) #create app instance
    root.mainloop() #continue till closed
